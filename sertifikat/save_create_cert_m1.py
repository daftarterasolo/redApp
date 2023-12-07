#!C:\Users\ibm\AppData\Local\Programs\Python\Python39\python.exe
import cgi
import cgitb
import json
import math
import statistics as stat
from datetime import datetime

# import library utk membuat file doc 
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches

# Import library utk membuat file excel
import xlsxwriter

# Import doc to pdf converter
from docx2pdf import convert

cgitb.enable()
print("Content-Type: text/html\n")
# print("Content-Type: application/json\n")

class hari_ini:
	def __init__(self):
		self.today = datetime.now().strftime("[%Y_%m_%d_%H_%M]")

	def tampilkan(self):
		return self.today

form = cgi.FieldStorage()
datas = form.getvalue("datas")
datas = json.loads(datas)
now_date = hari_ini().tampilkan()
# file_name = now_date + "[M1_5kg]"
no_surat = form.getvalue("no_surat")

massa_uji = []
for nilai in datas:
	massa_uji.append(str(int(nilai['unsur_data_pengujian']['at_uji']['m1_mass'])/1000)+"_kg")

# massa_uji = str(int(datas[0]['unsur_data_pengujian']['at_uji']['m1_mass'])/1000)
# file_name = now_date + "[M1_" + "_".join(massa_uji) + "kg]"
# file_name = now_date + "[M1_" + "_".join(massa_uji) + "]"
file_name = now_date + "[Kelas_M1]"

with open('dbase_json/massa/AT/M1/' + file_name + '.json', 'w') as testfile :
	json.dump(datas, testfile, indent=4,sort_keys=True)
	

#membuka file jason untuk membuat sertifikat
df = open('dbase_json/massa/AT/M1/' + file_name + '.json', 'r')	
expense_ = json.load(df)


#Membuat dokumen serifikat
doc = docx.Document('master_digital_signature.docx')

tgl_uji = expense_[0]['unsur_data_pengujian']['tgl_pengujian'].split()
obj_month = {'januari' : 'I', 'februari' : 'II', 'maret' : 'III', 'april' : 'IV', 'mei' : 'V', 'juni' : 'VI', 'juli' : 'VII', 'agustus' : 'VIII', 'september' : 'IX', 'oktober' : 'X', 'november' : 'XI', 'nopember' : 'XI', 'desember' : 'XII'}

doc.paragraphs[2].runs[1].text = 'DG.02.14 / '+str(no_surat)+' / '+obj_month[tgl_uji[1].lower().replace("'","")]+' / '+tgl_uji[2]

x = expense_[0]['unsur_data_pengujian']
y = expense_[0]['unsur_std']
z = expense_[0]['unsur_env']
doc.tables[0].rows[0].cells[3].text = x['nama_alat'].upper()
doc.tables[0].rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[0].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'

arr_of_merk = []

for k in expense_:
	arr_of_merk.append(k['unsur_data_pengujian']['data_penimbangan']['trademark'])

if all(elem == arr_of_merk[0] for elem in arr_of_merk):
	# doc.tables[0].rows[1].cells[3].text = arr_of_merk[0] +" ( 1 set )"
	doc.tables[0].rows[1].cells[3].text = arr_of_merk[0] +" / "+ k['unsur_data_pengujian']['data_penimbangan']['buatan']
else:
	doc.tables[0].rows[1].cells[3].text = '[ Terlampir ]'
	
doc.tables[0].rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[1].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'

arr_of_serial = []

for l in expense_:
	arr_of_serial.append(l['unsur_data_pengujian']['data_penimbangan']['s_1'])

if all(elem == arr_of_serial[0] for elem in arr_of_serial):
	# doc.tables[0].rows[2].cells[3].text = arr_of_serial[0] +" ( 1 set ) "
	doc.tables[0].rows[2].cells[3].text = arr_of_serial[0]
else:
	doc.tables[0].rows[2].cells[3].text = '[ Terlampir ]'
doc.tables[0].rows[2].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[2].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'

arr_of_massa = []

for i in expense_:
	arr_of_massa.append(i['unsur_data_pengujian']['at_uji']['m1_mass'])

# mengubah isi array arr_of_massa menjadi integer
arr_of_massa = [int(item) for item in arr_of_massa]
# mengurutkan isi arr_of_massa dari kecil ke besar
arr_of_massa.sort()

if all(elem == arr_of_massa[0] for elem in arr_of_massa):
	doc.tables[0].rows[3].cells[3].text = str(arr_of_massa[0])+" g"
else :
	doc.tables[0].rows[3].cells[3].text = str(arr_of_massa[0])+" g - "+str(arr_of_massa[len(arr_of_massa)-1])+" g (1 set)"
doc.tables[0].rows[3].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[3].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[0].rows[4].cells[3].text = x['at_uji']['m1_kelas']
doc.tables[0].rows[4].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[4].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[0].rows[6].cells[3].text = x['tgl_pengujian']
doc.tables[0].rows[6].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[6].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
# doc.tables[0].rows[9].cells[4].paragraphs[0].text = 'Surakarta, '+tgl_uji[0]+' '+tgl_uji[1]+' '+tgl_uji[2]
doc.tables[0].rows[9].cells[4].paragraphs[0].text = 'Surakarta, '+x['tgl_pengujian']
doc.tables[0].rows[9].cells[4].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[0].rows[9].cells[4].paragraphs[0].runs[0].font.name = 'Arial Narrow'

doc.tables[1].rows[1].cells[3].text = x['metode']
doc.tables[1].rows[1].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[1].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'

arr_of_std = []

for j in expense_:
	i = j['unsur_std']['at_std']
	arr_of_std.append(i['m1_kelas_reff']+" / "+i['m1_merk_at_ref']+" : "+i['m1_no_seri_atref'])

doc.tables[1].rows[2].cells[3].text = '['+']  ['.join(arr_of_std)+']'
doc.tables[1].rows[2].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[2].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[1].rows[3].cells[3].text = x['telusuran']
doc.tables[1].rows[3].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[3].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'

doc.tables[1].rows[6].cells[3].text = x['tgl_pengujian']
doc.tables[1].rows[6].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[6].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[1].rows[7].cells[3].text = x['penguji']['penguji_1'] +" / "+ x['penguji']['penguji_2']
doc.tables[1].rows[7].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[7].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[1].rows[8].cells[3].text = z['m1_lokasi']
doc.tables[1].rows[8].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[8].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[1].rows[9].cells[5].text = z['m1_temp']
doc.tables[1].rows[9].cells[5].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[9].cells[5].paragraphs[0].runs[0].font.name = 'Arial Narrow'
doc.tables[1].rows[10].cells[5].text = z['m1_kelembaban']
doc.tables[1].rows[10].cells[5].paragraphs[0].runs[0].font.size = Pt(12)
doc.tables[1].rows[10].cells[5].paragraphs[0].runs[0].font.name = 'Arial Narrow'


# hasil_tabel = doc.add_table(rows=len(expense_)+1,cols=6)
hasil_tabel = doc.add_table(rows=len(expense_)+1,cols=7)
hasil_tabel.style = 'Table Grid'
hdr_cells = hasil_tabel.rows[0].cells

# hdr_name = ['Merk', 'No. Seri', 'Massa Nominal (g)', 'Massa Konvensional (g)', 'U (g)', 'Status']
hdr_name = ['Merk', 'No. Seri', 'Massa Nominal (g)', 'Massa Konvensional (g)', 'U (g)', 'Bentuk AT', 'Status']

b = 0
for item in hdr_cells:
    item.text = hdr_name[b]
    b += 1
    item.paragraphs[0].runs[0].font.size = Pt(12)
    item.paragraphs[0].runs[0].font.name = 'Arial Narrow'
    item.paragraphs[0].runs[0].font.bold = True
    item.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

content_rows = hasil_tabel.rows


c=1
for x in expense_:
	input_U = x['unsur_data_pengujian']['hasil_pengujian']['U_luas']
	U = round(input_U, 2 - int(math.floor(math.log10(abs(input_U)))) - 1)
	tmp = str(U).split(".")
	len_U = len(tmp[1])
	content_rows[c].cells[0].text = x['unsur_data_pengujian']['data_penimbangan']['trademark']
	content_rows[c].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[0].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
	content_rows[c].cells[1].text = x['unsur_data_pengujian']['data_penimbangan']['s_1']
	content_rows[c].cells[1].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[1].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	content_rows[c].cells[2].text = str(round(int(x['unsur_data_pengujian']['at_uji']['m1_mass']), len_U))
	content_rows[c].cells[2].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[2].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
	
	content_rows[c].cells[3].text = str(round(x['unsur_data_pengujian']['hasil_pengujian']['mc_at_uji'], len_U))
	content_rows[c].cells[3].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[3].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
	content_rows[c].cells[4].text = str(U)
	content_rows[c].cells[4].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[4].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

	content_rows[c].cells[5].text = x['unsur_data_pengujian']['at_uji']['bentuk_at']
	content_rows[c].cells[5].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[5].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

	content_rows[c].cells[6].text = x['unsur_data_pengujian']['hasil_pengujian']['mc_sah']
	content_rows[c].cells[6].paragraphs[0].runs[0].font.size = Pt(12)
	content_rows[c].cells[6].paragraphs[0].runs[0].font.name = 'Arial Narrow'
	content_rows[c].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

	c += 1
	
doc.save('dbase_json/massa/AT/M1/'+ file_name +'[sertifikat].docx')
convert('dbase_json/massa/AT/M1/'+ file_name +'[sertifikat].docx')

# membuat file cerapan format excel


workbook = xlsxwriter.Workbook('dbase_json/massa/AT/M1/'+ file_name +'[cerapan].xlsx')

for i,val in enumerate(expense_):
	SheetName = expense_[i]['unsur_data_pengujian']['at_uji']['m1_mass']+' ke-'+str(i+1)
	worksheet = workbook.add_worksheet(SheetName)

	row = 1
	col = 0
	merge_format = workbook.add_format({
		'bold': 1,
		'border': 1,
		'align': 'center',
		'valign': 'vcenter',
		'fg_color':'#52B2BF'})
		
	cell_format = workbook.add_format({
		'border':1
	})

	cell_format2 = workbook.add_format({
		'border': 1,
		'bold': 1,
		'fg_color':'#FDA50F'
	})    

	cell_format_ttd_bold_underlined = workbook.add_format({
		'border': 0,
		'bold': 1,
		'underline' : 1
	})    

	cell_format_ttd = workbook.add_format({
		'border': 0,
		'bold': 0,
		'underline' : 0
	})    

	# worksheet.merge_range('A1:B1', 'PERSYARATAN', merge_format)
	
	name_arr = ['Massa Nominal', 'Kelas', 'BKD', 'Massa Jenis']   
	name_key = ['m1_mass','m1_kelas','m1_bkd','m1_density_t']
	satuan_arr = ['g','','g','kg/m3'] 
	row = 1
	col = 0
	worksheet.merge_range('A1:C1', 'DATA AT UJI', merge_format)

	_j_ = 0
	for j in name_arr:
		worksheet.write(row, col, name_arr[_j_], cell_format)
		worksheet.write(row, col+1, expense_[i]['unsur_data_pengujian']['at_uji'][name_key[_j_]], cell_format)
		worksheet.write(row, col+2, satuan_arr[_j_], cell_format)
		row += 1
		_j_ += 1

	
	name_arr = ['Merk', 'No Seri', 'Massa konvensional', 'U', 'Massa Jenis', 'Kelas', 'BKD']   
	satuan_arr = ['', '', 'g', 'g', 'kg/m3', '', 'g'] 
	name_key = ['m1_merk_at_ref','m1_no_seri_atref','m1_mass_conv','m1_u', 'm1_density_r', 'm1_kelas_reff', 'm1_bkd_reff']
	row = 1
	col = 4
	worksheet.merge_range('E1:G1', 'DATA AT STANDAR', merge_format)
	_j_ = 0
	for j in name_arr:
		worksheet.write(row, col, name_arr[_j_], cell_format)
		worksheet.write(row, col+1, expense_[i]['unsur_std']['at_std'][name_key[_j_]], cell_format)
		worksheet.write(row, col+2, satuan_arr[_j_], cell_format)
		row += 1
		_j_ += 1


	satuan_arr = ['', '', 'g', 'g', 'g']
	name_key = ['m1_merk_com','m1_serial_com','m1_kap_com','m1_d_com','m1_u_com'] 
	name_arr = ['Merk', 'Serial', 'Kapasitas', 'Dayabaca', 'U']
	row = 9
	col = 0
	worksheet.merge_range('A9:C9', 'DATA MASS COMPARATOR', merge_format)
	_j_ = 0

	for j in satuan_arr:
		worksheet.write(row, col+2, satuan_arr[_j_], cell_format)
		worksheet.write(row, col+1, expense_[i]['unsur_std']['masscom'][name_key[_j_]], cell_format)
		worksheet.write(row, col, name_arr[_j_], cell_format)
		row += 1
		_j_ += 1

	
	name_arr = ['Lokasi', 'Temperatur', 'Kelembaban Udara', 'Tekanan Udara', 'Massa Jenis Udara'] 
	satuan_arr = ['', ' derajat Celcius', '%', 'hPa', 'kg/m3'] 
	name_key = ['m1_lokasi','m1_temp','m1_kelembaban','m1_tekanan_u','m1_air_density']
	row = 9
	col = 4
	worksheet.merge_range('E9:G9', 'DATA LINGKUNGAN', merge_format)
	_j_ = 0
	for j in name_arr:
		worksheet.write(row, col, name_arr[_j_], cell_format)
		worksheet.write(row, col+1, expense_[i]['unsur_env'][name_key[_j_]], cell_format)
		worksheet.write(row, col+2, satuan_arr[_j_], cell_format)
		row += 1
		_j_ += 1   	

	row = 16
	col = 0
	worksheet.merge_range('A16:B16', 'Merk AT Uji', merge_format)
	worksheet.merge_range('C16:D16', 'No Seri AT Uji', merge_format)
	worksheet.merge_range('A17:B17', expense_[i]['unsur_data_pengujian']['data_penimbangan']['trademark'], cell_format)
	worksheet.merge_range('C17:D17', expense_[i]['unsur_data_pengujian']['data_penimbangan']['s_1'], cell_format)
	worksheet.write('A18', 'A(g)', merge_format)
	worksheet.write('B18', 'B(g)', merge_format)
	worksheet.write('C18', 'B(g)', merge_format)
	worksheet.write('D18', 'A(g)', merge_format)
	worksheet.write('A19', expense_[i]['unsur_data_pengujian']['data_penimbangan']['A1'], cell_format)
	worksheet.write('B19', expense_[i]['unsur_data_pengujian']['data_penimbangan']['B1'], cell_format)
	worksheet.write('C19', expense_[i]['unsur_data_pengujian']['data_penimbangan']['B_1'], cell_format)
	worksheet.write('D19', expense_[i]['unsur_data_pengujian']['data_penimbangan']['A_1'], cell_format)
	worksheet.write('E19', 'Seri-1', cell_format)

	worksheet.write('A20', expense_[i]['unsur_data_pengujian']['data_penimbangan']['A2'], cell_format)
	worksheet.write('B20', expense_[i]['unsur_data_pengujian']['data_penimbangan']['B2'], cell_format)
	worksheet.write('C20', expense_[i]['unsur_data_pengujian']['data_penimbangan']['B_2'], cell_format)
	worksheet.write('D20', expense_[i]['unsur_data_pengujian']['data_penimbangan']['A_2'], cell_format)
	worksheet.write('E20', 'Seri-2', cell_format)

	worksheet.write('A21', expense_[i]['unsur_data_pengujian']['data_penimbangan']['A3'], cell_format)
	worksheet.write('B21', expense_[i]['unsur_data_pengujian']['data_penimbangan']['B3'], cell_format)
	worksheet.write('C21', expense_[i]['unsur_data_pengujian']['data_penimbangan']['B_3'], cell_format)
	worksheet.write('D21', expense_[i]['unsur_data_pengujian']['data_penimbangan']['A_3'], cell_format)
	worksheet.write('E21', 'Seri-3', cell_format)
	worksheet.write('A22', 'mc', merge_format)
	worksheet.write('A23', 'selisih', merge_format)
	worksheet.write('B22', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['mc_at_uji'], cell_format)
	worksheet.write('B23', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['selisih'], cell_format)
	worksheet.write('A24', 'Status', merge_format)
	worksheet.write('B24', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['mc_sah'], cell_format)

	row = 26
	col = 0
	n = 0
	worksheet.merge_range('A26:B26', 'DRIFT AT STANDAR', merge_format)
	arr_drift_name = ['mc thn-3', 'mc thn-2', 'mc thn-1']

	if expense_[i]['unsur_data_pengujian']['at_uji']['m1_mass'] != '25000':
		name_key = ['drift_3','drift_2','drift_1']
		
		for drift in arr_drift_name:
			worksheet.write(row, col, arr_drift_name[n], cell_format2)
			worksheet.write(row, col+1, expense_[i]['unsur_std']['drift'][name_key[n]], cell_format)
			worksheet.write(row, col+2, '( g )')
			n += 1
			row += 1
		
	else :
		name_key = ['drift_3','drift_3_5kg','drift_2','drift_2_5kg','drift_1']
		
		for drift in arr_drift_name:
			
			if n == 2:
				sum = float(expense_[i]['unsur_std']['drift'][name_key[2*n]])
			else:
				sum = float(expense_[i]['unsur_std']['drift'][name_key[2*n]]) + float(expense_[0]['unsur_std']['drift'][name_key[(2*n)+1]])
				
			worksheet.write(row, col, arr_drift_name[n], cell_format2)
			worksheet.write(row, col+1, sum, cell_format)
			worksheet.write(row, col+2, '( g )')
			n += 1
			row += 1

	row = 32
	col = 0
	y = 0
	arr_U = ['Ketidaktetapan Penimbangan', 'Ketidakpastian AT Standar', 'Daya Baca Timbangan', 'Koreksi Bouyancy Udara', 'Drift AT Standar', 'Hitung Ketidakpastian', 'Uc (Ketidakpastian Gabungan)', 'Veff (Derajat Kebebasan Effektif)', 'U (Ketidakpastian Yang Diperluas', 'status'];
	worksheet.merge_range('A31:G31', 'PERHITUNGAN KETIDAKPASTIAN', merge_format)
	worksheet.write('A32', 'Komponen Ketidakpastian', merge_format)
	worksheet.write('B32', 'u', merge_format)
	worksheet.write('C32', 'c', merge_format)
	worksheet.write('D32', 'v', merge_format)
	worksheet.write('E32', '(u.c)2', merge_format)
	worksheet.write('F32', '(u.c)4', merge_format)
	worksheet.write('G32', '(u.c)4/v', merge_format)

	keys_horizontal = ['u','c','v','uc2','uc4','uc4v']
	keys_horizontal2 = ['Udb','cdb','vdb','U_c_2','U_c_4','U_c_4_v']
	keys_vertical = ['penimbangan','at_standar','d_mass_com','bouyancy','drift_at_std']


	def horizontal_fill(fill_arr, key_vertical, startRow, startCol, cell_format):
		x = 0
		col = startCol
		for j in fill_arr:
			worksheet.write(startRow, col, expense_[i]['unsur_ketidakpastian'][key_vertical][fill_arr[x]], cell_format)
			col += 1
			x += 1
		
	for T in arr_U:
		if (y <= 4):
			worksheet.write(row, 0, arr_U[y], cell_format)

			if (keys_vertical[y] == 'drift_at_std' or keys_vertical[y] == 'penimbangan' ):
				horizontal_fill(keys_horizontal, keys_vertical[y], row, 1, cell_format)
			else:
				horizontal_fill(keys_horizontal2, keys_vertical[y], row, 1, cell_format)
	
		if (y == 5):
			worksheet.merge_range('A38:D38', 'J U M L A H : ', merge_format)
			worksheet.write_formula('E38', '=sum(E49:E45)', merge_format)
			worksheet.write_formula('F38', '=sum(F49:F45)', merge_format)
			worksheet.write_formula('G38', '=sum(G49:G45)', merge_format)
		
		if (y == 6):
			worksheet.merge_range('A39:D39', arr_U[y], merge_format)
			worksheet.write('E39',expense_[i]['unsur_data_pengujian']['hasil_pengujian']['U_gabungan'],cell_format)
		if (y == 7):
			worksheet.merge_range('A40:D40', arr_U[y], merge_format)
			worksheet.write('E40', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['V_eff'], cell_format)                
		if (y == 8):
			worksheet.merge_range('A41:D41', arr_U[y], merge_format)    
			worksheet.write('E41', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['U_luas'], cell_format)
		if (y == 9):
			worksheet.merge_range('F39:G41', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['u_sah'], merge_format)
		
		y += 1
		row += 1	

	if (expense_[i]['unsur_data_pengujian']['penguji']['penguji_1'] != ""):
		penguji_1 = expense_[i]['unsur_data_pengujian']['penguji']['penguji_1'].split(" - ")
	else :
		penguji_1 = " - "

	if (expense_[i]['unsur_data_pengujian']['penguji']['penguji_2'] != ""):
		penguji_2 = expense_[i]['unsur_data_pengujian']['penguji']['penguji_2'].split(" - ") 
	else :
		penguji_2 = " - "
		
	worksheet.write('A43', 'Penguji-1', cell_format_ttd)
	worksheet.write('F43', 'Penguji-2', cell_format_ttd)
	worksheet.write('A46', penguji_1[0], cell_format_ttd_bold_underlined)
	worksheet.write('F46', penguji_2[0], cell_format_ttd_bold_underlined)
	worksheet.write('A47', penguji_1[1], cell_format_ttd)
	worksheet.write('F47', penguji_2[1], cell_format_ttd)

"""
	row = 26
	col = 0
	n = 0
	worksheet.merge_range('A26:B26', 'DRIFT AT STANDAR', merge_format)
	arr_drift_name = ['mc thn-3', 'mc thn-2', 'mc thn-1']

	if expense_[i]['unsur_data_pengujian']['at_uji']['m1_mass'] != '25000':
		name_key = ['drift_3','drift_2','drift_1']
		
		for drift in arr_drift_name:
			worksheet.write(row, col, arr_drift_name[n], cell_format2)
			worksheet.write(row, col+1, expense_[i]['unsur_std']['drift'][name_key[n]], cell_format)
			worksheet.write(row, col+2, '( g )')
			n += 1
			row += 1
		
	else :
		name_key = ['drift_3','drift_3_5kg','drift_2','drift_2_5kg','drift_1']
		
		for drift in arr_drift_name:
			
			if n == 2:
				sum = float(expense_[i]['unsur_std']['drift'][name_key[2*n]])
			else:
				sum = float(expense_[i]['unsur_std']['drift'][name_key[2*n]]) + float(expense_[0]['unsur_std']['drift'][name_key[(2*n)+1]])
				
			worksheet.write(row, col, arr_drift_name[n], cell_format2)
			worksheet.write(row, col+1, sum, cell_format)
			worksheet.write(row, col+2, '( g )')
			n += 1
			row += 1
		
	row = 32
	col = 0
	y = 0
	arr_U = ['Ketidaktetapan Penimbangan', 'Ketidakpastian AT Standar', 'Daya Baca Timbangan', 'Koreksi Bouyancy Udara', 'Drift AT Standar', 'Hitung Ketidakpastian', 'Uc (Ketidakpastian Gabungan)', 'Veff (Derajat Kebebasan Effektif)', 'U (Ketidakpastian Yang Diperluas', 'status'];
	worksheet.merge_range('A31:G31', 'PERHITUNGAN KETIDAKPASTIAN', merge_format)
	worksheet.write('A32', 'Komponen Ketidakpastian', merge_format)
	worksheet.write('B32', 'u', merge_format)
	worksheet.write('C32', 'c', merge_format)
	worksheet.write('D32', 'v', merge_format)
	worksheet.write('E32', '(u.c)2', merge_format)
	worksheet.write('F32', '(u.c)4', merge_format)
	worksheet.write('G32', '(u.c)4/v', merge_format)


	keys_horizontal = ['u','c','v','uc2','uc4','uc4v']
	keys_horizontal2 = ['Udb','cdb','vdb','U_c_2','U_c_4','U_c_4_v']
	keys_vertical = ['penimbangan','at_standar','d_mass_com','bouyancy','drift_at_std']


	def horizontal_fill(fill_arr, key_vertical, startRow, startCol, cell_format):
		x = 0
		col = startCol
		for j in fill_arr:
			worksheet.write(startRow, col, expense_[i]['unsur_ketidakpastian'][key_vertical][fill_arr[x]], cell_format)
			col += 1
			x += 1
		
	
	for T in arr_U:
		if (y <= 4):
			worksheet.write(row, 0, arr_U[y], cell_format)

			if (keys_vertical[y] == 'drift_at_std' or keys_vertical[y] == 'penimbangan' ):
				horizontal_fill(keys_horizontal, keys_vertical[y], row, 1, cell_format)
			else:
				horizontal_fill(keys_horizontal2, keys_vertical[y], row, 1, cell_format)
	
		if (y == 5):
			worksheet.merge_range('A38:D38', 'J U M L A H : ', merge_format)
			worksheet.write_formula('E38', '=sum(E49:E45)', merge_format)
			worksheet.write_formula('F38', '=sum(F49:F45)', merge_format)
			worksheet.write_formula('G38', '=sum(G49:G45)', merge_format)
		
		if (y == 6):
			worksheet.merge_range('A39:D39', arr_U[y], merge_format)
			worksheet.write('E39',expense_[i]['unsur_data_pengujian']['hasil_pengujian']['U_gabungan'],cell_format)
		if (y == 7):
			worksheet.merge_range('A40:D40', arr_U[y], merge_format)
			worksheet.write('E40', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['V_eff'], cell_format)                
		if (y == 8):
			worksheet.merge_range('A41:D41', arr_U[y], merge_format)    
			worksheet.write('E41', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['U_luas'], cell_format)
		if (y == 9):
			worksheet.merge_range('F39:G41', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['u_sah'], merge_format)
		
		y += 1
		row += 1	
	
	penguji_1 = expense_[i]['unsur_data_pengujian']['penguji']['penguji_1'].split(" - ") 
	penguji_2 = expense_[i]['unsur_data_pengujian']['penguji']['penguji_2'].split(" - ") 
	worksheet.write('A43', 'Penguji-1', cell_format_ttd)
	worksheet.write('F43', 'Penguji-2', cell_format_ttd)
	worksheet.write('A46', penguji_1[0], cell_format_ttd_bold_underlined)
	worksheet.write('F46', penguji_2[0], cell_format_ttd_bold_underlined)
	worksheet.write('A47', penguji_1[1], cell_format_ttd)
	worksheet.write('F47', penguji_2[1], cell_format_ttd)
"""	
	
"""	
	row = 38
	col = 0
	n = 0
	worksheet.merge_range('A38:B38', 'DRIFT AT STANDAR', merge_format)
	arr_drift_name = ['mc thn-3', 'mc thn-2', 'mc thn-1']

	if expense_[i]['unsur_data_pengujian']['at_uji']['m1_mass'] != '25000':
		name_key = ['drift_3','drift_2','drift_1']
		
		for drift in arr_drift_name:
			worksheet.write(row, col, arr_drift_name[n], cell_format2)
			worksheet.write(row, col+1, expense_[i]['unsur_std']['drift'][name_key[n]], cell_format)
			worksheet.write(row, col+2, '( g )')
			n += 1
			row += 1
		
	else :
		name_key = ['drift_3','drift_3_5kg','drift_2','drift_2_5kg','drift_1']
		
		for drift in arr_drift_name:
			
			if n == 2:
				sum = float(expense_[i]['unsur_std']['drift'][name_key[2*n]])
			else:
				sum = float(expense_[i]['unsur_std']['drift'][name_key[2*n]]) + float(expense_[0]['unsur_std']['drift'][name_key[(2*n)+1]])
				
			worksheet.write(row, col, arr_drift_name[n], cell_format2)
			worksheet.write(row, col+1, sum, cell_format)
			worksheet.write(row, col+2, '( g )')
			n += 1
			row += 1
		
	row = 44
	col = 0
	y = 0
	arr_U = ['Ketidaktetapan Penimbangan', 'Ketidakpastian AT Standar', 'Daya Baca Timbangan', 'Koreksi Bouyancy Udara', 'Drift AT Standar', 'Hitung Ketidakpastian', 'Uc (Ketidakpastian Gabungan)', 'Veff (Derajat Kebebasan Effektif)', 'U (Ketidakpastian Yang Diperluas', 'status'];
	worksheet.merge_range('A43:G43', 'PERHITUNGAN KETIDAKPASTIAN', merge_format)
	worksheet.write('A44', 'Komponen Ketidakpastian', merge_format)
	worksheet.write('B44', 'u', merge_format)
	worksheet.write('C44', 'c', merge_format)
	worksheet.write('D44', 'v', merge_format)
	worksheet.write('E44', '(u.c)2', merge_format)
	worksheet.write('F44', '(u.c)4', merge_format)
	worksheet.write('G44', '(u.c)4/v', merge_format)


	keys_horizontal = ['u','c','v','uc2','uc4','uc4v']
	keys_horizontal2 = ['Udb','cdb','vdb','U_c_2','U_c_4','U_c_4_v']
	keys_vertical = ['penimbangan','at_standar','d_mass_com','bouyancy','drift_at_std']


	def horizontal_fill(fill_arr, key_vertical, startRow, startCol, cell_format):
		x = 0
		col = startCol
		for j in fill_arr:
			worksheet.write(startRow, col, expense_[i]['unsur_ketidakpastian'][key_vertical][fill_arr[x]], cell_format)
			col += 1
			x += 1
		
	
	for T in arr_U:
		if (y <= 4):
			worksheet.write(row, 0, arr_U[y], cell_format)

			if (keys_vertical[y] == 'drift_at_std' or keys_vertical[y] == 'penimbangan' ):
				horizontal_fill(keys_horizontal, keys_vertical[y], row, 1, cell_format)
			else:
				horizontal_fill(keys_horizontal2, keys_vertical[y], row, 1, cell_format)
	
		if (y == 5):
			worksheet.merge_range('A50:D50', 'J U M L A H : ', merge_format)
			worksheet.write_formula('E50', '=sum(E49:E45)', merge_format)
			worksheet.write_formula('F50', '=sum(F49:F45)', merge_format)
			worksheet.write_formula('G50', '=sum(G49:G45)', merge_format)
		
		if (y == 6):
			worksheet.merge_range('A51:D51', arr_U[y], merge_format)
			worksheet.write('E51',expense_[i]['unsur_data_pengujian']['hasil_pengujian']['U_gabungan'],cell_format)
		if (y == 7):
			worksheet.merge_range('A52:D52', arr_U[y], merge_format)
			worksheet.write('E52', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['V_eff'], cell_format)                
		if (y == 8):
			worksheet.merge_range('A53:D53', arr_U[y], merge_format)    
			worksheet.write('E53', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['U_luas'], cell_format)
		if (y == 9):
			worksheet.merge_range('F51:G53', expense_[i]['unsur_data_pengujian']['hasil_pengujian']['u_sah'], merge_format)
		
		y += 1
		row += 1	
"""								
workbook.close() 

"""
list_of_files = {"file_cert" : 'dbase_json/massa/AT/M2/certdata/'+now_date+'[sertifikat].pdf', "file_cerapan" :'dbase_json/massa/AT/M2/certdata/'+now_date+'[cerapan].xlsx', "file_json" : 'dbase_json/massa/AT/M2/certdata/'+now_date+'[sertifikat].json'}
"""


list_of_files = {"file_cert" : 'dbase_json/massa/AT/M1/'+ file_name +'[sertifikat].pdf', "file_cerapan" :'dbase_json/massa/AT/M1/'+ file_name +'[cerapan].xlsx', "file_json" : 'dbase_json/massa/AT/M1/'+ file_name +'.json'}
print(json.dumps(list_of_files))
