#/home/metskasrv/.nix-profile/htdocs/redApp/sertifikat/sertifikat-env/bin/python3
import cgi
import cgitb
import json
import math
import statistics as stat
from datetime import datetime

# import library utk membuat file doc 
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches

# Import library utk membuat file excel
import xlsxwriter

# Import doc to pdf converter
from docx2pdf import convert

cgitb.enable()
print("Content-Type: text/html\n")

doc = docx.Document('master.docx')


"""
fulltext = []
k = 0
for para in doc.paragraphs:
	for run in para:
		print(str(k)+run)
"""

"""
for baris in doc.tables[1]:
	print(baris)	 
"""

print(doc.tables[0].rows[0].cells[0].text)
#print(doc.paragraphs[6].runs[1].text)
#for tbl in doc.tables:
	#print(tbl.rows[0].cells[0].text)	
#print(doc.paragraphs[6].runs[0].text)
#print(doc.paragraphs[6].runs[1].text)
#print(doc.paragraphs[6].runs[2].text)
#print(fulltext[2])
